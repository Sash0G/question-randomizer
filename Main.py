import random
import openpyxl

def read_excel_file(file_path):
    wb = openpyxl.load_workbook(file_path)
    sheet = wb.active
    
    data = []
    for row in sheet.iter_rows(values_only=True):
       if row[0] is not None: data.append(row)
    wb.close
    return data
    

from docx import Document

def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)
p = input()
for k in range(int(p)):
    doc = Document()
    formatted_data = ""
    questions = []
    answers = []
    for i in range(4):
        file_path = "set"+str(i+1)+".xlsx"
        # print(file_path)
        set = read_excel_file(file_path)
        # print(set)
        random.shuffle(set)
        for j in range(20):
            try: questions.append(set[j][0])
            except: break
            answers.append(set[j][1])

    for num,item in enumerate(questions):
        doc.add_paragraph(str(num+1)+". "+item)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'  # Example table with 1 row and 3 columns
    for j,row_data in enumerate(answers):
        row_cells = table.rows[-1].cells
        row_cells[0].text = str(j+1)
        row_cells[1].text = str(row_data)
        if j < len(answers)-1: table.add_row()
    # Print or wr.ite the formatted data

    # Or write to a file
    doc.save("Test"+str(k+1)+".docx")

